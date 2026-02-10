
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def markdown_to_docx(md_path, output_path):
    doc = Document()

    # Define styles
    styles = doc.styles
    
    # Title Style
    if 'Title' in styles:
        title_style = styles['Title']
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(26)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 51, 102) # Dark Blue

    # Heading 1
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 51, 102)

    # Normal text
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)

    # Add Logo if exists
    logo_path = r"C:/Users/Marco/.gemini/antigravity/brain/cebdbbb7-ba52-4d53-9f51-90b8729b31eb/techsos_logo_v7_fixed_handle_1769175263786.png"
    if os.path.exists(logo_path):



        doc.add_picture(logo_path, width=Inches(2.5))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []

    for line in lines:
        line = line.strip()

        # Skip separators
        if line == '---':
            doc.add_page_break()
            continue
        
        # Handle Table
        if line.startswith('|'):
            in_table = True
            table_lines.append(line)
            continue
        else:
            if in_table:
                # Process the accumulated table
                process_table(doc, table_lines)
                table_lines = []
                in_table = False
            
            if not line:
                continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            runner = p.add_run(line[2:])
            runner.italic = True
            runner.font.color.rgb = RGBColor(80, 80, 80)
            p.paragraph_format.left_indent = Inches(0.5)
        elif line.startswith('* '):
             p = doc.add_paragraph(style='List Bullet')
             process_bold_text(p, line[2:])
        elif line.startswith('1. '):
             p = doc.add_paragraph(style='List Number')
             process_bold_text(p, line[3:])
        else:
            p = doc.add_paragraph()
            process_bold_text(p, line)

    # If file ends with table
    if in_table:
        process_table(doc, table_lines)

    doc.save(output_path)
    print(f"Document saved to {output_path}")

def process_bold_text(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 == 1: # Odd parts are inside ** **
            run.bold = True

def process_table(doc, lines):
    # Filter header separator line like |---|---|
    data_lines = [l for l in lines if '---' not in l]
    
    if not data_lines:
        return

    # Determine columns
    first_line = data_lines[0].strip('|').split('|')
    cols = len(first_line)
    
    table = doc.add_table(rows=len(data_lines), cols=cols)
    table.style = 'Table Grid'
    
    for i, line in enumerate(data_lines):
        cells = line.strip('|').split('|')
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(cells):
            if j < len(row_cells):
                p = row_cells[j].paragraphs[0]
                process_bold_text(p, cell_text.strip())
                if i == 0: # Header
                     for run in p.runs:
                         run.bold = True
                         run.font.color.rgb = RGBColor(0, 0, 0)
                         # Simple shading isn't easy in pure python-docx without Oxml hack, skip for now

if __name__ == "__main__":
    md_file = r"c:\Users\Marco\.gemini\antigravity\brain\cebdbbb7-ba52-4d53-9f51-90b8729b31eb\business_plan.md"
    docx_file = r"c:\Users\Marco\.gemini\antigravity\brain\cebdbbb7-ba52-4d53-9f51-90b8729b31eb\Plano_de_Negocios_TechSOS.docx"
    markdown_to_docx(md_file, docx_file)
